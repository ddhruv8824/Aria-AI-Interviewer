"""
services/resume/parser.py
──────────────────────────
ATS Resume Parser
=================
Extracts structured data from ATS-formatted resumes (PDF and DOCX).

Requirements:
    pip install pdfplumber python-docx

Usage (standalone):
    python -m services.resume.parser resume.pdf
    python -m services.resume.parser resume.docx
"""

import re
import sys
import json
import argparse
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Optional

try:
    from google import genai
    from google.genai import types
except ImportError:
    genai = None
    types = None

try:
    from shared.config import API_KEY
except ImportError:
    API_KEY = None


# ── Data model ────────────────────────────────────────────────────────────────

@dataclass
class ContactInfo:
    name:     Optional[str] = None
    email:    Optional[str] = None
    phone:    Optional[str] = None
    linkedin: Optional[str] = None
    github:   Optional[str] = None
    location: Optional[str] = None
    website:  Optional[str] = None


@dataclass
class Experience:
    company:    Optional[str] = None
    title:      Optional[str] = None
    start_date: Optional[str] = None
    end_date:   Optional[str] = None
    location:   Optional[str] = None
    bullets:    list[str] = field(default_factory=list)


@dataclass
class Education:
    institution:     Optional[str] = None
    degree:          Optional[str] = None
    field_of_study:  Optional[str] = None
    graduation_date: Optional[str] = None
    gpa:             Optional[str] = None


@dataclass
class ResumeData:
    source_file:    str = ""
    contact:        ContactInfo = field(default_factory=ContactInfo)
    summary:        Optional[str] = None
    skills:         list[str] = field(default_factory=list)
    experience:     list[Experience] = field(default_factory=list)
    education:      list[Education] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    languages:      list[str] = field(default_factory=list)
    raw_text:       str = ""


# ── Text extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(path: Path) -> str:
    """Extract raw text from a PDF using pdfplumber (preserves layout)."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Run: pip install pdfplumber")

    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                lines.append(text)
    return "\n".join(lines)


def extract_text_from_docx(path: Path) -> str:
    """Extract raw text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Run: pip install python-docx")

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF or DOCX.")


# ── Regex patterns ─────────────────────────────────────────────────────────────

EMAIL_RE    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE    = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")
LINKEDIN_RE = re.compile(r"(linkedin\.com/in/[\w\-]+)", re.IGNORECASE)
GITHUB_RE   = re.compile(r"(github\.com/[\w\-]+)", re.IGNORECASE)
URL_RE      = re.compile(r"https?://[\w./\-?=#%&]+", re.IGNORECASE)
GPA_RE      = re.compile(r"GPA[:\s]*([\d.]+)", re.IGNORECASE)
DATE_RE     = re.compile(
    r"(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s*\d{4}|"
    r"\d{1,2}/\d{4}|\d{4}",
    re.IGNORECASE,
)
DATE_RANGE_RE = re.compile(
    r"("
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s*\d{4}|\d{1,2}/\d{4}|\d{4}"
    r")"
    r"\s*[-–—to]+\s*"
    r"("
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s*\d{4}|\d{1,2}/\d{4}|\d{4}|[Pp]resent|[Cc]urrent"
    r")",
    re.IGNORECASE,
)

SECTION_HEADERS = {
    "summary":        re.compile(r"^\s*(summary|profile|objective|about me)\s*$", re.I),
    "experience":     re.compile(r"^\s*(experience|work experience|employment|professional experience|work history)\s*$", re.I),
    "education":      re.compile(r"^\s*(education|academic background|qualifications)\s*$", re.I),
    "skills":         re.compile(r"^\s*(skills|technical skills|core competencies|competencies|technologies)\s*$", re.I),
    "certifications": re.compile(r"^\s*(certifications?|licenses?|credentials)\s*$", re.I),
    "languages":      re.compile(r"^\s*(languages?)\s*$", re.I),
}

DEGREE_KEYWORDS = re.compile(
    r"(B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|MBA|PhD|Ph\.D|Associate|Bachelor|Master|Doctor)",
    re.IGNORECASE,
)


# ── Parsing helpers ────────────────────────────────────────────────────────────

def detect_section(line: str) -> Optional[str]:
    for section, pattern in SECTION_HEADERS.items():
        if pattern.match(line):
            return section
    return None


def parse_contact(lines: list[str], contact: ContactInfo) -> None:
    """Scan the top portion of the resume for contact fields."""
    top = lines[:30]
    full_text_top = "\n".join(top)

    email = EMAIL_RE.search(full_text_top)
    if email:
        contact.email = email.group()

    phone = PHONE_RE.search(full_text_top)
    if phone:
        raw = phone.group().strip()
        contact.phone = re.sub(r"[^\d+\s\-()]", "", raw).strip()

    linkedin = LINKEDIN_RE.search(full_text_top)
    if linkedin:
        contact.linkedin = "https://" + linkedin.group()

    github = GITHUB_RE.search(full_text_top)
    if github:
        contact.github = "https://" + github.group()

    urls = URL_RE.findall(full_text_top)
    for url in urls:
        if "linkedin" not in url and "github" not in url:
            contact.website = url
            break

    for line in lines[:8]:
        stripped = line.strip()
        words = stripped.split()
        if (
            2 <= len(words) <= 5
            and not any(ch.isdigit() for ch in stripped)
            and "@" not in stripped
        ):
            contact.name = stripped
            break


def split_into_sections(text: str) -> dict[str, list[str]]:
    """Split resume text into named sections."""
    sections: dict[str, list[str]] = {"header": []}
    current = "header"

    for line in text.splitlines():
        detected = detect_section(line)
        if detected:
            current = detected
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)

    return sections


def parse_skills(lines: list[str]) -> list[str]:
    skills = []
    for line in lines:
        line = line.strip().strip("•·-–—▪►*")
        if not line:
            continue
        parts = re.split(r"[,|;/]", line)
        for part in parts:
            part = part.strip()
            if part and len(part) < 60:
                skills.append(part)
    return skills


def parse_experience(lines: list[str]) -> list[Experience]:
    entries: list[Experience] = []
    current: Optional[Experience] = None

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        date_match = DATE_RANGE_RE.search(line_stripped)
        if date_match:
            if current:
                entries.append(current)
            current = Experience()
            current.start_date = date_match.group(1)
            current.end_date   = date_match.group(2)
            before_date = line_stripped[: date_match.start()].strip(" |–-—·,")
            if before_date:
                parts = re.split(r"\s{2,}|[|,]", before_date, maxsplit=1)
                current.title = parts[0].strip() if parts else before_date
                if len(parts) > 1:
                    current.company = parts[1].strip()
        elif current:
            bullet = line_stripped.lstrip("•·-–—▪►*● ").strip()
            if bullet:
                current.bullets.append(bullet)

    if current:
        entries.append(current)
    return entries


def parse_education(lines: list[str]) -> list[Education]:
    entries: list[Education] = []
    current: Optional[Education] = None

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        degree_match = DEGREE_KEYWORDS.search(line_stripped)
        gpa_match    = GPA_RE.search(line_stripped)
        date_match   = DATE_RE.search(line_stripped)

        if degree_match:
            if current:
                entries.append(current)
            current = Education()
            current.degree = line_stripped
            if date_match:
                current.graduation_date = date_match.group()
        elif current:
            if gpa_match:
                current.gpa = gpa_match.group(1)
            elif date_match and not current.graduation_date:
                current.graduation_date = date_match.group()
            elif not current.institution and line_stripped:
                current.institution = line_stripped

    if current:
        entries.append(current)
    return entries


def parse_certifications(lines: list[str]) -> list[str]:
    certs = []
    for line in lines:
        line = line.strip().lstrip("•·-–—▪►*● ")
        if line:
            certs.append(line)
    return certs


def parse_languages(lines: list[str]) -> list[str]:
    langs = []
    for line in lines:
        line = line.strip().lstrip("•·-–—▪►*● ")
        parts = re.split(r"[,|;]", line)
        langs.extend(p.strip() for p in parts if p.strip())
    return langs


# ── Main parser ────────────────────────────────────────────────────────────────

def _parse_resume_with_gemini(path: Path) -> Optional[dict]:
    if not genai or not API_KEY:
        return None
    try:
        client = genai.Client(api_key=API_KEY)
        
        prompt = """
        You are an expert ATS (Applicant Tracking System) resume parser.
        Analyze the uploaded resume (PDF or DOCX or raw text) and extract all information accurately.
        
        Specifically ensure you read multi-column layouts correctly. Do not interlace lines or text across columns.
        
        You must output a single JSON object matching this schema exactly:
        {
          "contact": {
            "name": "Full Name",
            "email": "Email Address",
            "phone": "Phone Number",
            "linkedin": "LinkedIn URL (or null)",
            "github": "GitHub URL (or null)",
            "location": "City, State/Country (or null)",
            "website": "Personal Website URL (or null)"
          },
          "summary": "Short professional summary or objective, if present (or null)",
          "skills": ["Skill 1", "Skill 2", ...],
          "experience": [
            {
              "company": "Company Name",
              "title": "Job Title",
              "start_date": "Start Date (e.g. Month Year)",
              "end_date": "End Date or 'Present'",
              "location": "Location (City, State/Country or null)",
              "bullets": [
                "Responsibility / achievement bullet 1",
                "Responsibility / achievement bullet 2"
              ]
            }
          ],
          "education": [
            {
              "institution": "University/School Name",
              "degree": "Degree (e.g. B.S., BCA)",
              "field_of_study": "Field of Study (or null)",
              "graduation_date": "Graduation Date/Year (e.g. 2021 - 2024)",
              "gpa": "GPA (if present or null)"
            }
          ],
          "certifications": ["Certification 1", "Certification 2", ...],
          "languages": ["Language 1", "Language 2", ...]
        }
        """
        
        ext = path.suffix.lower()
        if ext == ".pdf":
            # Pass PDF bytes directly for layout preservation
            pdf_bytes = path.read_bytes()
            contents = [
                types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"),
                prompt
            ]
        else:
            # Fallback to extracted text for other formats or if binary read fails
            raw_text = extract_text(path)
            contents = [
                f"Here is the raw text of the resume:\n\n{raw_text}",
                prompt
            ]
            
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=contents,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.1
            )
        )
        
        return json.loads(response.text.strip())
    except Exception as exc:
        print(f"Warning: Gemini resume parser failed, falling back to regex parser: {exc}", file=sys.stderr)
        return None

def parse_resume(path: Path, original_filename: Optional[str] = None) -> ResumeData:
    gemini_data = _parse_resume_with_gemini(path)
    
    raw_text = ""
    try:
        raw_text = extract_text(path)
    except Exception:
        pass

    if gemini_data:
        try:
            # Reconstruct contact
            contact_dict = gemini_data.get("contact", {})
            contact = ContactInfo(
                name=contact_dict.get("name"),
                email=contact_dict.get("email"),
                phone=contact_dict.get("phone"),
                linkedin=contact_dict.get("linkedin"),
                github=contact_dict.get("github"),
                location=contact_dict.get("location"),
                website=contact_dict.get("website"),
            )
            
            # Reconstruct experience
            experience = []
            for exp in gemini_data.get("experience", []):
                experience.append(Experience(
                    company=exp.get("company"),
                    title=exp.get("title"),
                    start_date=exp.get("start_date"),
                    end_date=exp.get("end_date"),
                    location=exp.get("location"),
                    bullets=exp.get("bullets") or [],
                ))
                
            # Reconstruct education
            education = []
            for edu in gemini_data.get("education", []):
                education.append(Education(
                    institution=edu.get("institution"),
                    degree=edu.get("degree"),
                    field_of_study=edu.get("field_of_study"),
                    graduation_date=edu.get("graduation_date"),
                    gpa=edu.get("gpa"),
                ))
                
            data = ResumeData(
                source_file=str(path),
                contact=contact,
                summary=gemini_data.get("summary"),
                skills=gemini_data.get("skills") or [],
                experience=experience,
                education=education,
                certifications=gemini_data.get("certifications") or [],
                languages=gemini_data.get("languages") or [],
                raw_text=raw_text
            )
        except Exception as e:
            print(f"Warning: Failed to construct ResumeData from Gemini response: {e}. Falling back to regex parser.", file=sys.stderr)
            gemini_data = None
            
    if not gemini_data:
        # Fallback to regex-based parser
        lines    = raw_text.splitlines()
        sections = split_into_sections(raw_text)

        data = ResumeData(source_file=str(path), raw_text=raw_text)

        parse_contact(lines, data.contact)

        summary_lines = sections.get("summary", [])
        summary_text  = " ".join(l.strip() for l in summary_lines if l.strip())
        if summary_text:
            data.summary = summary_text

        data.skills         = parse_skills(sections.get("skills", []))
        data.experience     = parse_experience(sections.get("experience", []))
        data.education      = parse_education(sections.get("education", []))
        data.certifications = parse_certifications(sections.get("certifications", []))
        data.languages      = parse_languages(sections.get("languages", []))

    # Save parsed data to 'parsed_resumes' folder for validation
    try:
        project_root = Path(__file__).resolve().parent.parent.parent.parent
        output_dir = project_root / "parsed_resumes"
        output_dir.mkdir(parents=True, exist_ok=True)

        fname = original_filename or path.name
        base_fname = Path(fname).stem
        clean_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', base_fname)

        import datetime
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = output_dir / f"{clean_name}_{timestamp}.json"

        serialized = asdict(data)
        serialized["parsed_at"] = datetime.datetime.now().isoformat()
        serialized["parser_used"] = "gemini" if gemini_data else "regex_fallback"

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(serialized, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error saving parsed resume to JSON: {e}", file=sys.stderr)

    return data


# ── CLI (standalone usage) ─────────────────────────────────────────────────────

def _pretty_print(data: ResumeData) -> None:
    c = data.contact
    print("\n" + "=" * 60)
    print(f"  RESUME: {Path(data.source_file).name}")
    print("=" * 60)
    print(f"\n📋 CONTACT")
    print(f"  Name      : {c.name or '—'}")
    print(f"  Email     : {c.email or '—'}")
    print(f"  Phone     : {c.phone or '—'}")
    print(f"  LinkedIn  : {c.linkedin or '—'}")
    print(f"  GitHub    : {c.github or '—'}")

    if data.skills:
        print(f"\n🛠  SKILLS ({len(data.skills)})")
        print("  " + ", ".join(data.skills[:20]))

    if data.experience:
        print(f"\n💼 EXPERIENCE ({len(data.experience)} roles)")
        for exp in data.experience:
            print(f"\n  {exp.title or 'Unknown'} @ {exp.company or 'Unknown'}")
            print(f"  {exp.start_date or '?'} – {exp.end_date or '?'}")
            for b in exp.bullets[:3]:
                print(f"    • {b[:100]}")

    if data.education:
        print(f"\n🎓 EDUCATION ({len(data.education)} entries)")
        for edu in data.education:
            print(f"  {edu.degree or '—'}")
            print(f"  {edu.institution or '—'}  |  Graduated: {edu.graduation_date or '—'}")

    if data.certifications:
        print(f"\n🏅 CERTIFICATIONS ({len(data.certifications)})")
        for cert in data.certifications:
            print(f"  • {cert}")

    if data.languages:
        print(f"\n🌐 LANGUAGES")
        print("  " + ", ".join(data.languages))
    print()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="ATS Resume Parser")
    parser.add_argument("path", help="Path to a resume file (PDF/DOCX) or directory")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    parser.add_argument("--out",  help="Save JSON output to this file")
    args = parser.parse_args()

    input_path = Path(args.path)
    results = []

    files = (
        list(input_path.glob("*.pdf")) + list(input_path.glob("*.docx"))
        if input_path.is_dir()
        else [input_path]
    )
    if not files:
        print("No PDF or DOCX files found.")
        sys.exit(1)

    for f in files:
        print(f"Parsing: {f.name} ...")
        try:
            data = parse_resume(f)
            results.append(data)
            if not args.json:
                _pretty_print(data)
        except Exception as e:
            print(f"  ⚠ Error: {e}")

    if args.json or args.out:
        output = json.dumps([asdict(r) for r in results], indent=2, ensure_ascii=False)
        if args.out:
            Path(args.out).write_text(output, encoding="utf-8")
            print(f"\n✅ JSON saved to {args.out}")
        else:
            print(output)
